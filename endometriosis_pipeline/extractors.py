"""Streaming extractors for corpus archives and common document formats."""

from __future__ import annotations

import csv
import html
import io
import json
import mimetypes
from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document

from .ocr import AnthropicOCR


TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".log",
    ".text",
}
IMAGE_SUFFIXES = {
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".webp",
    ".tif",
    ".tiff",
    ".bmp",
}
SUPPORTED_SUFFIXES = (
    TEXT_SUFFIXES
    | IMAGE_SUFFIXES
    | {".csv", ".tsv", ".json", ".jsonl", ".ndjson", ".zst", ".pdf", ".docx", ".html", ".htm"}
)
CONTENT_KEYS = {
    "title",
    "selftext",
    "body",
    "text",
    "content",
    "description",
    "caption",
    "summary",
    "abstract",
    "message",
}
METADATA_KEYS = {
    "id",
    "author",
    "username",
    "subreddit",
    "permalink",
    "url",
    "parent_id",
    "created_utc",
    "timestamp",
}


@dataclass
class TextUnit:
    unit: str
    text: str


@dataclass
class SourceExtractor:
    method: str
    units: Iterator[TextUnit]


class UnsupportedFileError(ValueError):
    pass


def discover_files(input_dir: Path) -> list[Path]:
    return sorted(
        path
        for path in input_dir.rglob("*")
        if path.is_file() and not path.name.startswith(".")
    )


def _clean(value: str) -> str:
    return html.unescape(value).replace("\x00", "").strip()


def _strings_from_json(value: object) -> Iterator[str]:
    if isinstance(value, str):
        cleaned = _clean(value)
        if cleaned:
            yield cleaned
        return

    if isinstance(value, list):
        for item in value:
            yield from _strings_from_json(item)
        return

    if not isinstance(value, dict):
        return

    content_found = False
    for key, item in value.items():
        if str(key).lower() in CONTENT_KEYS:
            content_found = True
            yield from _strings_from_json(item)

    for key, item in value.items():
        normalized_key = str(key).lower()
        if normalized_key in CONTENT_KEYS or normalized_key in METADATA_KEYS:
            continue
        if isinstance(item, (dict, list)) or not content_found:
            yield from _strings_from_json(item)


def _text_file(path: Path) -> Iterator[TextUnit]:
    with path.open("r", encoding="utf-8", errors="replace") as handle:
        buffer: list[str] = []
        for line_number, line in enumerate(handle, 1):
            buffer.append(line)
            if len(buffer) >= 500:
                yield TextUnit(f"lines {line_number - len(buffer) + 1}-{line_number}", "".join(buffer))
                buffer = []
        if buffer:
            end = line_number if "line_number" in locals() else len(buffer)
            yield TextUnit(f"lines {end - len(buffer) + 1}-{end}", "".join(buffer))


def _delimited_file(path: Path) -> Iterator[TextUnit]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for row_number, row in enumerate(reader, 1):
            text = "\t".join(_clean(cell) for cell in row if _clean(cell))
            if text:
                yield TextUnit(f"row {row_number}", text)


def _json_file(path: Path) -> Iterator[TextUnit]:
    with path.open("r", encoding="utf-8", errors="replace") as handle:
        value = json.load(handle)
    for index, text in enumerate(_strings_from_json(value), 1):
        yield TextUnit(f"value {index}", text)


def _json_lines(handle: io.TextIOBase) -> Iterator[TextUnit]:
    for line_number, line in enumerate(handle, 1):
        line = line.strip()
        if not line:
            continue
        try:
            value = json.loads(line)
            parts = list(_strings_from_json(value))
            text = "\n".join(parts)
        except json.JSONDecodeError:
            text = _clean(line)
        if text:
            yield TextUnit(f"line {line_number}", text)


def _jsonl_file(path: Path) -> Iterator[TextUnit]:
    with path.open("r", encoding="utf-8", errors="replace") as handle:
        yield from _json_lines(handle)


def _zst_file(path: Path) -> Iterator[TextUnit]:
    import zstandard as zstd

    with path.open("rb") as compressed:
        decompressor = zstd.ZstdDecompressor(max_window_size=2**31)
        with decompressor.stream_reader(compressed) as reader:
            text_reader = io.TextIOWrapper(reader, encoding="utf-8", errors="replace")
            yield from _json_lines(text_reader)


def _docx_file(path: Path) -> Iterator[TextUnit]:
    document = Document(path)
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = _clean(paragraph.text)
        if text:
            yield TextUnit(f"paragraph {index}", text)
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = "\t".join(_clean(cell.text) for cell in row.cells if _clean(cell.text))
            if text:
                yield TextUnit(f"table {table_index}, row {row_index}", text)


def _html_file(path: Path) -> Iterator[TextUnit]:
    markup = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(markup, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = "\n".join(
        line.strip() for line in soup.get_text("\n").splitlines() if line.strip()
    )
    if text:
        yield TextUnit("document", text)


def _pdf_file(
    path: Path,
    ocr: AnthropicOCR | None,
    force_ocr: bool,
) -> Iterator[TextUnit]:
    import fitz

    with fitz.open(path) as document:
        for page_number, page in enumerate(document, 1):
            embedded = _clean(page.get_text("text"))
            needs_ocr = force_ocr or len("".join(embedded.split())) < 40
            if needs_ocr and ocr is not None:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                text = ocr.transcribe_bytes(pixmap.tobytes("png"), "image/png")
            else:
                text = embedded
            if text:
                yield TextUnit(f"page {page_number}", text)


def extractor_for(
    path: Path,
    ocr: AnthropicOCR | None = None,
    force_ocr: bool = False,
) -> SourceExtractor:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return SourceExtractor("plain text", _text_file(path))
    if suffix in {".csv", ".tsv"}:
        return SourceExtractor("delimited text", _delimited_file(path))
    if suffix == ".json":
        return SourceExtractor("JSON text fields", _json_file(path))
    if suffix in {".jsonl", ".ndjson"}:
        return SourceExtractor("JSON Lines text fields", _jsonl_file(path))
    if suffix == ".zst":
        return SourceExtractor("Zstandard JSON Lines text fields", _zst_file(path))
    if suffix == ".docx":
        return SourceExtractor("DOCX text", _docx_file(path))
    if suffix in {".html", ".htm"}:
        return SourceExtractor("HTML visible text", _html_file(path))
    if suffix == ".pdf":
        method = "PDF text with OCR fallback" if ocr else "embedded PDF text"
        return SourceExtractor(method, _pdf_file(path, ocr, force_ocr))
    if suffix in IMAGE_SUFFIXES:
        if ocr is None:
            raise UnsupportedFileError(
                f"{path.name} needs OCR, but OCR is disabled or no API key is configured"
            )
        return SourceExtractor(
            "Anthropic image OCR",
            iter([TextUnit("image", ocr.transcribe_file(path))]),
        )
    guessed_type = mimetypes.guess_type(path.name)[0] or "unknown"
    raise UnsupportedFileError(f"unsupported file type: {suffix or guessed_type}")
